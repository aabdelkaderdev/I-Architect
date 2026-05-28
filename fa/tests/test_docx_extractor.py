import pytest
from docx import Document
from ingestion.extractors import extract_from_docx
from ingestion.exceptions import ExtractionError, FormatMismatchError
import os

def test_extract_from_docx_paragraphs(tmp_path):
    doc = Document()
    doc.add_paragraph("This is a valid paragraph with more than 10 characters.")
    doc.add_paragraph("Too short")
    doc.add_paragraph("   ")
    
    file_path = tmp_path / "test.docx"
    doc.save(str(file_path))
    
    blocks = extract_from_docx(str(file_path))
    assert len(blocks) == 1
    assert blocks[0]["text"] == "This is a valid paragraph with more than 10 characters."
    assert blocks[0]["source_section"] is None

def test_extract_from_docx_lists(tmp_path):
    doc = Document()
    p = doc.add_paragraph("1. First list item", style="List Number")
    p = doc.add_paragraph("• Second list item", style="List Bullet")
    
    file_path = tmp_path / "test_lists.docx"
    doc.save(str(file_path))
    
    blocks = extract_from_docx(str(file_path))
    assert len(blocks) == 2
    assert blocks[0]["text"] == "First list item"
    assert blocks[1]["text"] == "Second list item"

def test_extract_from_docx_sections(tmp_path):
    doc = Document()
    doc.add_heading("1.0 System Requirements", level=1)
    doc.add_paragraph("This paragraph belongs to the system requirements section.")
    doc.add_heading("2.0 Unrelated Header", level=1)
    doc.add_paragraph("This paragraph has no recognized section.")
    
    file_path = tmp_path / "test_sections.docx"
    doc.save(str(file_path))
    
    blocks = extract_from_docx(str(file_path))
    assert len(blocks) == 2
    assert blocks[0]["text"] == "This paragraph belongs to the system requirements section."
    assert blocks[0]["source_section"] == "1.0 System Requirements"
    assert blocks[1]["text"] == "This paragraph has no recognized section."
    assert blocks[1]["source_section"] is None

def test_extract_from_docx_tables(tmp_path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "ID"
    table.rows[0].cells[1].text = "Description"
    table.rows[1].cells[0].text = "REQ-1"
    table.rows[1].cells[1].text = "System shall do X."
    
    file_path = tmp_path / "test_table.docx"
    doc.save(str(file_path))
    
    blocks = extract_from_docx(str(file_path))
    assert len(blocks) == 1
    assert blocks[0]["text"] == "ID: REQ-1 | Description: System shall do X."

def test_extract_from_docx_empty(tmp_path):
    doc = Document()
    file_path = tmp_path / "empty.docx"
    doc.save(str(file_path))
    
    with pytest.raises(ExtractionError, match="No text content found"):
        extract_from_docx(str(file_path))

def test_extract_from_docx_corrupt(tmp_path):
    file_path = tmp_path / "corrupt.docx"
    file_path.write_text("not a docx")
    
    with pytest.raises(FormatMismatchError):
        extract_from_docx(str(file_path))
