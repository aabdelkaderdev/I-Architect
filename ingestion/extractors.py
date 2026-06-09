import json
import os
import pdfplumber
import docx
import chardet
from typing import Any
import collections
import re
import unicodedata

from ingestion.exceptions import (
    EmptyFileError,
    UnsupportedFormatError,
    ExtractionError,
    NonStandardJSONError,
    FormatMismatchError
)

def validate_json_schema(data: Any) -> None:
    if not isinstance(data, dict):
        raise NonStandardJSONError("Root value is not a dict")
        
    offending_keys = []
    for k, v in data.items():
        # Check if key is a non-empty string
        if not isinstance(k, str) or not k.strip():
            offending_keys.append(str(k))
            continue
            
        # Check if value is a non-empty string
        if not isinstance(v, str) or not v.strip():
            offending_keys.append(str(k))
            continue
            
    if offending_keys:
        raise NonStandardJSONError("Keys and values must be non-empty strings and not contain nested structures", offending_keys)

def extract_from_json(file_path: str) -> dict[str, str]:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        raise FormatMismatchError(f"File cannot be parsed as JSON: {e}")
        
    validate_json_schema(data)
    return data

def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = "".join(c for c in text if c == '\n' or c == '\t' or ord(c) >= 0x20)
    text = unicodedata.normalize("NFKD", text)
    text = text.replace('\u2018', "'").replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"')
    return text

def _is_semantic_header(header: str) -> bool:
    if not header:
        return False
    keywords = {"id", "requirement", "description", "shall", "must", "functional", "non-functional", "constraint"}
    lower_header = header.lower()
    return any(kw in lower_header for kw in keywords)

def _process_table(table: list[list[str]], page_num: int) -> list[dict]:
    if not table or len(table) < 2:
        return []
    
    headers = [str(c).strip() if c else "" for c in table[0]]
    semantic = any(_is_semantic_header(h) for h in headers)
    
    table_blocks = []
    for row in table[1:]:
        cells = [str(c).strip() if c else "" for c in row]
        if semantic:
            parts = [f"{h}: {c}" if h else c for h, c in zip(headers, cells) if c]
            block_text = " | ".join(parts)
        else:
            block_text = " ".join(c for c in cells if c)
            
        block_text = _clean_text(block_text)
        if block_text and block_text.strip():
            table_blocks.append({
                "text": block_text.strip(),
                "source_page": page_num,
                "source_section": None
            })
    return table_blocks

def _segment_text(text: str, page_num: int) -> list[dict]:
    blocks = []
    list_pattern = re.compile(r'^\s*(?:\d+[\.\)]\s+|[a-zA-Z][\.\)]\s+|•\s+|[-–]\s+)', re.MULTILINE)
    matches = list(list_pattern.finditer(text))
    
    if matches:
        for i in range(len(matches)):
            match_start = matches[i].start()
            if i == 0 and match_start > 0:
                pre_text = text[0:match_start].strip()
                if pre_text:
                    blocks.append(pre_text)
            next_match_start = matches[i+1].start() if i + 1 < len(matches) else len(text)
            item_text = text[match_start:next_match_start].strip()
            if item_text:
                blocks.append(item_text)
    else:
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
        for p in paragraphs:
            if len(p) > 500:
                sentence_pattern = re.compile(r'(?<=\.)\s+(?=[A-Z])')
                sentences = [s.strip() for s in sentence_pattern.split(p) if s.strip()]
                blocks.extend(sentences)
            else:
                blocks.append(p)
                
    results = []
    for b in blocks:
        clean_b = _clean_text(b)
        if clean_b and clean_b.strip():
            results.append({
                "text": clean_b.strip(),
                "source_page": page_num,
                "source_section": None
            })
    return results

def extract_from_pdf(file_path: str, pdf_engine: str = "pdfplumber", header_footer_threshold: float = 0.6) -> list[dict]:
    extracted_tables = []
    pages_text = []
    has_any_text = False
    
    try:
        if pdf_engine == "pdfplumber":
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_num = i + 1
                    tables = page.find_tables()
                    for table in tables:
                        raw_table = table.extract()
                        if raw_table:
                            extracted_tables.extend(_process_table(raw_table, page_num))
                    
                    if tables:
                        bboxes = [t.bbox for t in tables]
                        def not_in_bboxes(obj):
                            if 'x0' not in obj or 'top' not in obj or 'x1' not in obj or 'bottom' not in obj:
                                return True
                            center_x = (obj['x0'] + obj['x1']) / 2
                            center_y = (obj['top'] + obj['bottom']) / 2
                            for (x0, top, x1, bottom) in bboxes:
                                if x0 <= center_x <= x1 and top <= center_y <= bottom:
                                    return False
                            return True
                        page_text_obj = page.filter(not_in_bboxes)
                        page_text = page_text_obj.extract_text()
                    else:
                        page_text = page.extract_text()
                        
                    if page_text:
                        has_any_text = True
                        pages_text.append((page_num, page_text))
                    else:
                        pages_text.append((page_num, ""))

        elif pdf_engine == "pymupdf":
            import fitz
            with fitz.open(file_path) as doc:
                total_pages = len(doc)
                for i, page in enumerate(doc):
                    page_num = i + 1
                    if hasattr(page, 'find_tables'):
                        for tab in page.find_tables():
                            raw_table = tab.extract()
                            if raw_table:
                                extracted_tables.extend(_process_table(raw_table, page_num))
                    
                    page_text = page.get_text()
                    if page_text:
                        has_any_text = True
                        pages_text.append((page_num, page_text))
                    else:
                        pages_text.append((page_num, ""))
        else:
            raise ValueError(f"Unknown PDF engine: {pdf_engine}")
    except Exception as e:
        raise FormatMismatchError(f"Could not parse as PDF: {e}")
        
    if not has_any_text:
        raise ExtractionError("No extractable text; OCR not supported")
        
    line_counts = collections.Counter()
    page_lines = []
    
    for page_num, text in pages_text:
        lines = text.split('\n')
        norm_lines = []
        for line in lines:
            norm = " ".join(line.split())
            if norm:
                norm_lines.append(norm)
        
        for unique_line in set(norm_lines):
            line_counts[unique_line] += 1
            
        page_lines.append((page_num, lines))
        
    threshold_count = total_pages * header_footer_threshold
    headers_footers = {line for line, count in line_counts.items() if count > threshold_count}
    
    final_blocks = []
    final_blocks.extend(extracted_tables)
    
    for page_num, lines in page_lines:
        filtered_lines = []
        for line in lines:
            norm = " ".join(line.split())
            if norm and norm not in headers_footers:
                filtered_lines.append(line)
                
        filtered_text = "\n".join(filtered_lines)
        if filtered_text.strip():
            final_blocks.extend(_segment_text(filtered_text, page_num))
            
    return final_blocks

def extract_from_docx(file_path: str) -> list[dict]:
    blocks = []
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        raise FormatMismatchError(f"Could not parse as DOCX: {e}")
        
    current_section = None
    current_heading_level = 99
    
    list_marker_regex = re.compile(r'^\s*(?:\d+[\.\)]\s+|[a-zA-Z][\.\)]\s+|•\s+|[-–]\s+)')
    
    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        
        # Heading tracking
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except ValueError:
                level = 1
                
            heading_keywords = {"requirements", "functional requirements", "non-functional requirements", "business requirements", "system requirements"}
            if any(kw in text.lower() for kw in heading_keywords):
                current_section = text
                current_heading_level = level
            elif level <= current_heading_level:
                # Reset if we hit a higher or equal heading without keywords
                current_section = None
                current_heading_level = 99
            # Headings are not emitted as blocks
            continue
            
        if "List" in style_name:
            text = list_marker_regex.sub('', text).strip()
            
        if len(text) < 10:
            continue
                
        blocks.append({
            "text": text,
            "source_page": None,
            "source_section": current_section
        })
        
    # Process tables
    for table in doc.tables:
        if not table.rows:
            continue
            
        # Empty table check (all cells whitespace)
        is_empty = True
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    is_empty = False
                    break
            if not is_empty:
                break
        if is_empty:
            continue
            
        headers = [c.text.strip() for c in table.rows[0].cells]
        semantic_count = sum(1 for h in headers if _is_semantic_header(h))
        semantic = semantic_count >= 2
        
        start_idx = 1 if semantic else 0
        
        for row in table.rows[start_idx:]:
            cells = [c.text.strip() for c in row.cells]
            
            if semantic:
                parts = []
                for i in range(min(len(headers), len(cells))):
                    if cells[i]:
                        h = headers[i]
                        if h:
                            parts.append(f"{h}: {cells[i]}")
                        else:
                            parts.append(cells[i])
                for i in range(len(headers), len(cells)):
                    if cells[i]:
                        parts.append(cells[i])
                block_text = " | ".join(parts)
            else:
                block_text = " ".join(c for c in cells if c)
                
            block_text = _clean_text(block_text)
            if block_text and block_text.strip():
                blocks.append({
                    "text": block_text.strip(),
                    "source_page": None,
                    "source_section": None
                })
                
    if not blocks:
        raise ExtractionError("No text content found")
        
    return blocks

def extract_from_txt(file_path: str, encoding_fallback: str = "utf-8") -> list[dict]:
    with open(file_path, 'rb') as f:
        raw_data = f.read()
        
    if not raw_data.strip():
        raise ExtractionError("No text content found")
        
    result = chardet.detect(raw_data)
    confidence = result['confidence'] or 0.0
    encoding = result['encoding'] if confidence >= 0.7 and result['encoding'] else encoding_fallback
    
    try:
        text = raw_data.decode(encoding)
    except Exception:
        try:
            text = raw_data.decode(encoding_fallback)
        except Exception as e:
            raise ExtractionError(f"Could not decode text file with detected ({encoding}) or fallback ({encoding_fallback}) encoding: {e}")
            
    if not text.strip():
        raise ExtractionError("No text content found")
        
    lines = [line.strip() for line in text.splitlines()]
    non_empty_lines = [line for line in lines if line]
    if not non_empty_lines:
        raise ExtractionError("No text content found")
        
    structured_pattern = re.compile(r'^([A-Za-z]+-?\d+)\s*[:\.]\s*(.+)$')
    structured_matches = [structured_pattern.match(line) for line in non_empty_lines]
    structured_count = sum(1 for m in structured_matches if m)
    
    blocks = []
    list_marker_regex = re.compile(r'^\s*(?:\d+[\.\)]\s+|[a-zA-Z][\.\)]\s+|•\s+|[-–]\s+)')
    
    if len(non_empty_lines) > 0 and (structured_count / len(non_empty_lines) > 0.5):
        # Structured mode
        for line in non_empty_lines:
            match = structured_pattern.match(line)
            if match:
                inline_id = match.group(1).strip()
                cleaned_line = list_marker_regex.sub('', match.group(2)).strip()
                if cleaned_line:
                    blocks.append({
                        "text": cleaned_line,
                        "inline_id": inline_id,
                        "source_page": None,
                        "source_section": None
                    })
    else:
        # Unstructured mode
        avg_len = sum(len(line) for line in non_empty_lines) / len(non_empty_lines) if non_empty_lines else 0
        if avg_len < 200:
            # Line mode
            for line in non_empty_lines:
                cleaned = list_marker_regex.sub('', line).strip()
                if cleaned:
                    blocks.append({
                        "text": cleaned,
                        "source_page": None,
                        "source_section": None
                    })
        else:
            # Paragraph mode
            paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
            for p in paragraphs:
                p_flat = " ".join(p.split())
                cleaned = list_marker_regex.sub('', p_flat).strip()
                if cleaned:
                    blocks.append({
                        "text": cleaned,
                        "source_page": None,
                        "source_section": None
                    })
                    
    if not blocks:
        raise ExtractionError("No text content found")
        
    return blocks
